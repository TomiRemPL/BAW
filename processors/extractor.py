"""Document extraction module for DOCX files."""
import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

import docx2python
from docx import Document
from pydantic import BaseModel

from exceptions import DocumentExtractionError, EmptyDocumentError


logger = logging.getLogger(__name__)


class DocumentMetadata(BaseModel):
    """Metadata extracted from a document."""
    author: Optional[str] = None
    last_modified: Optional[datetime] = None
    page_count: int = 0
    paragraph_count: int = 0
    table_count: int = 0


class TableStructure(BaseModel):
    """Structure representing a table."""
    index: int
    rows: List[List[str]]
    row_count: int
    col_count: int


class Comment(BaseModel):
    """Structure representing a comment."""
    author: str
    text: str
    date: Optional[datetime] = None


class ExtractedContent(BaseModel):
    """Complete extracted content from a document."""
    metadata: DocumentMetadata
    paragraphs: List[str]
    tables: List[TableStructure]
    comments: List[Comment]
    raw_text: str


class DocumentExtractor:
    """Extracts content from DOCX files."""

    def __init__(self, verbose: bool = True):
        """
        Initialize the DocumentExtractor.

        Args:
            verbose: Enable verbose logging
        """
        self.verbose = verbose
        if verbose:
            logger.setLevel(logging.DEBUG)

    def extract(self, file_path: str) -> ExtractedContent:
        """
        Extract all content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            ExtractedContent with all extracted data

        Raises:
            DocumentExtractionError: If extraction fails
            EmptyDocumentError: If document has no content
        """
        path = Path(file_path)

        logger.info(f"Extracting content from: {file_path}")

        try:
            # Extract metadata and structured content using python-docx
            metadata = self._extract_metadata(file_path)
            paragraphs = self._extract_paragraphs(file_path)
            tables = self._extract_tables(file_path)
            comments = self._extract_comments(file_path)

            # Extract raw text using docx2python for full text
            raw_text = self._extract_raw_text(file_path)

            # Validate that we got some content
            if not paragraphs and not tables and not raw_text.strip():
                logger.error(f"Document appears to be empty: {file_path}")
                raise EmptyDocumentError(
                    f"Dokument nie zawiera żadnej treści: {path.name}"
                )

            logger.info(
                f"Extraction complete: {len(paragraphs)} paragraphs, "
                f"{len(tables)} tables, {len(comments)} comments"
            )

            return ExtractedContent(
                metadata=metadata,
                paragraphs=paragraphs,
                tables=tables,
                comments=comments,
                raw_text=raw_text
            )

        except EmptyDocumentError:
            raise

        except Exception as e:
            logger.error(f"Error extracting document: {e}", exc_info=True)
            raise DocumentExtractionError(
                f"Nie udało się wyodrębnić dokumentu {path.name}: {str(e)}"
            )

    def _extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract metadata from document."""
        logger.debug("Extracting metadata")

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            # Count elements
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            metadata = DocumentMetadata(
                author=core_props.author,
                last_modified=core_props.modified,
                page_count=0,  # Not easily available in python-docx
                paragraph_count=paragraph_count,
                table_count=table_count
            )

            logger.debug(f"Metadata extracted: {metadata}")
            return metadata

        except Exception as e:
            logger.warning(f"Error extracting metadata: {e}")
            return DocumentMetadata()

    def _extract_paragraphs(self, file_path: str) -> List[str]:
        """Extract paragraphs preserving structure."""
        logger.debug("Extracting paragraphs")

        try:
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)

            logger.debug(f"Extracted {len(paragraphs)} paragraphs")
            return paragraphs

        except Exception as e:
            logger.error(f"Error extracting paragraphs: {e}")
            return []

    def _extract_tables(self, file_path: str) -> List[TableStructure]:
        """Extract tables with structure."""
        logger.debug("Extracting tables")

        try:
            doc = Document(file_path)
            tables = []

            for idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)

                if rows:  # Only include non-empty tables
                    table_struct = TableStructure(
                        index=idx,
                        rows=rows,
                        row_count=len(rows),
                        col_count=len(rows[0]) if rows else 0
                    )
                    tables.append(table_struct)

            logger.debug(f"Extracted {len(tables)} tables")
            return tables

        except Exception as e:
            logger.error(f"Error extracting tables: {e}")
            return []

    def _extract_comments(self, file_path: str) -> List[Comment]:
        """Extract comments from document."""
        logger.debug("Extracting comments")

        # Note: python-docx doesn't have native comment support
        # This is a placeholder for future enhancement
        logger.debug("Comment extraction not fully implemented")
        return []

    def _extract_raw_text(self, file_path: str) -> str:
        """Extract raw text using docx2python."""
        logger.debug("Extracting raw text with docx2python")

        try:
            with docx2python.docx2python(file_path, html=True) as doc_result:
                # Get text content
                raw_text = doc_result.text

            logger.debug(f"Extracted {len(raw_text)} characters of raw text")
            return raw_text

        except Exception as e:
            logger.error(f"Error extracting raw text: {e}")
            return ""
