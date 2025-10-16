"""Moduł ekstrakcji treści z dokumentów DOCX."""
import logging
from pathlib import Path
from typing import List, Optional
from datetime import datetime

import docx2python
from docx import Document
from pydantic import BaseModel


logger = logging.getLogger(__name__)


class DocumentMetadata(BaseModel):
    """Metadane wyodrębnione z dokumentu."""
    author: Optional[str] = None
    last_modified: Optional[datetime] = None
    page_count: int = 0
    paragraph_count: int = 0
    table_count: int = 0


class TableStructure(BaseModel):
    """Struktura reprezentująca tabelę."""
    index: int
    rows: List[List[str]]
    row_count: int
    col_count: int


class ExtractedContent(BaseModel):
    """Kompletna wyodrębniona treść z dokumentu."""
    metadata: DocumentMetadata
    paragraphs: List[str]
    tables: List[TableStructure]
    raw_text: str


class DocumentExtractor:
    """Ekstraktor treści z plików DOCX."""

    def __init__(self, verbose: bool = False):
        """
        Inicjalizacja DocumentExtractor.

        Args:
            verbose: Włącz szczegółowe logowanie
        """
        self.verbose = verbose
        if verbose:
            logger.setLevel(logging.DEBUG)

    def extract(self, file_path: str) -> ExtractedContent:
        """
        Wyodrębnienie całej treści z pliku DOCX.

        Args:
            file_path: Ścieżka do pliku DOCX

        Returns:
            ExtractedContent ze wszystkimi wyodrębnionymi danymi

        Raises:
            Exception: Jeśli ekstrakcja się nie powiedzie
        """
        path = Path(file_path)
        logger.info(f"Ekstrakcja treści z: {file_path}")

        try:
            metadata = self._extract_metadata(file_path)
            paragraphs = self._extract_paragraphs(file_path)
            tables = self._extract_tables(file_path)
            raw_text = self._extract_raw_text(file_path)

            if not paragraphs and not tables and not raw_text.strip():
                raise Exception(f"Dokument nie zawiera żadnej treści: {path.name}")

            logger.info(
                f"Ekstrakcja zakończona: {len(paragraphs)} paragrafów, "
                f"{len(tables)} tabel"
            )

            return ExtractedContent(
                metadata=metadata,
                paragraphs=paragraphs,
                tables=tables,
                raw_text=raw_text
            )

        except Exception as e:
            logger.error(f"Błąd podczas ekstrakcji dokumentu: {e}", exc_info=True)
            raise Exception(f"Nie udało się wyodrębnić dokumentu {path.name}: {str(e)}")

    def _extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Wyodrębnienie metadanych z dokumentu."""
        logger.debug("Ekstrakcja metadanych")

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            metadata = DocumentMetadata(
                author=core_props.author,
                last_modified=core_props.modified,
                page_count=0,
                paragraph_count=paragraph_count,
                table_count=table_count
            )

            logger.debug(f"Metadane wyodrębnione: {metadata}")
            return metadata

        except Exception as e:
            logger.warning(f"Błąd podczas ekstrakcji metadanych: {e}")
            return DocumentMetadata()

    def _extract_paragraphs(self, file_path: str) -> List[str]:
        """Wyodrębnienie paragrafów z zachowaniem struktury."""
        logger.debug("Ekstrakcja paragrafów")

        try:
            doc = Document(file_path)
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            logger.debug(f"Wyodrębniono {len(paragraphs)} paragrafów")
            return paragraphs

        except Exception as e:
            logger.error(f"Błąd podczas ekstrakcji paragrafów: {e}")
            return []

    def _extract_tables(self, file_path: str) -> List[TableStructure]:
        """Wyodrębnienie tabel ze strukturą."""
        logger.debug("Ekstrakcja tabel")

        try:
            doc = Document(file_path)
            tables = []

            for idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(cells)

                if rows:
                    table_struct = TableStructure(
                        index=idx,
                        rows=rows,
                        row_count=len(rows),
                        col_count=len(rows[0]) if rows else 0
                    )
                    tables.append(table_struct)

            logger.debug(f"Wyodrębniono {len(tables)} tabel")
            return tables

        except Exception as e:
            logger.error(f"Błąd podczas ekstrakcji tabel: {e}")
            return []

    def _extract_raw_text(self, file_path: str) -> str:
        """Wyodrębnienie surowego tekstu używając docx2python."""
        logger.debug("Ekstrakcja surowego tekstu z docx2python")

        try:
            with docx2python.docx2python(file_path, html=False) as doc_result:
                raw_text = doc_result.text

            logger.debug(f"Wyodrębniono {len(raw_text)} znaków surowego tekstu")
            return raw_text

        except Exception as e:
            logger.error(f"Błąd podczas ekstrakcji surowego tekstu: {e}")
            return ""
