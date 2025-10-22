"""
Post-processing dla dokumentów prawnych/bankowych.

Poprawia formatowanie typowe dla dokumentów:
- Numeracja sekcji (1. 1.1. 1.1.1.)
- Definicje i klauzule
- Tabele finansowe
- Przypisy
"""

import logging
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


logger = logging.getLogger(__name__)


class PostProcessor:
    """
    Post-processor dla dokumentów prawnych i bankowych.

    Poprawia:
    - Numerację sekcji
    - Formatowanie list
    - Tabele
    - Definicje
    """

    def process(self, docx_path: Path) -> bool:
        """
        Przetwarza dokument DOCX.

        Args:
            docx_path: Ścieżka do DOCX

        Returns:
            True jeśli sukces, False w przeciwnym razie
        """
        try:
            logger.info(f"Post-processing: {docx_path}")

            doc = Document(docx_path)

            # 1. Popraw numerację
            self._fix_numbering(doc)

            # 2. Popraw tabele
            self._fix_tables(doc)

            # 3. Popraw definicje
            self._fix_definitions(doc)

            # 4. Popraw formatowanie list
            self._fix_lists(doc)

            # Zapisz
            doc.save(docx_path)

            logger.info(f"Post-processing zakończony: {docx_path}")
            return True

        except Exception as e:
            logger.error(f"Błąd post-processing: {e}")
            return False

    def _fix_numbering(self, doc: Document) -> None:
        """
        Poprawia numerację sekcji (1. 1.1. 1.1.1.).

        Args:
            doc: Dokument python-docx
        """
        # Pattern dla numeracji: 1., 1.1., 1.1.1., etc.
        numbering_pattern = re.compile(r'^(\d+\.)+\s+')

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Sprawdź czy to numerowana sekcja
            if numbering_pattern.match(text):
                # Określ poziom wcięcia na podstawie liczby kropek
                dots_count = text.count('.') - 1  # -1 bo ostatnia kropka to koniec

                # Ustaw wcięcie
                if dots_count == 0:
                    # Główny nagłówek (1.)
                    paragraph.style = 'Heading 1'
                elif dots_count == 1:
                    # Podsekcja (1.1.)
                    paragraph.style = 'Heading 2'
                elif dots_count == 2:
                    # Pod-podsekcja (1.1.1.)
                    paragraph.style = 'Heading 3'

    def _fix_tables(self, doc: Document) -> None:
        """
        Poprawia formatowanie tabel.

        Args:
            doc: Dokument python-docx
        """
        for table in doc.tables:
            # Ustaw styl tabeli
            table.style = 'Table Grid'

            # Pierwsza linia jako nagłówek (pogrubienie)
            if len(table.rows) > 0:
                header_row = table.rows[0]
                for cell in header_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

    def _fix_definitions(self, doc: Document) -> None:
        """
        Poprawia formatowanie definicji i klauzul.

        Typowy format: "Słowo" oznacza ... lub Definicja: ...

        Args:
            doc: Dokument python-docx
        """
        # Pattern dla definicji
        definition_pattern = re.compile(r'^"([^"]+)"\s+(oznacza|to|definiuje)', re.IGNORECASE)

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Sprawdź czy to definicja
            match = definition_pattern.match(text)
            if match:
                # Pogrub termin w cudzysłowie
                defined_term = match.group(1)
                for run in paragraph.runs:
                    if defined_term in run.text:
                        run.bold = True

    def _fix_lists(self, doc: Document) -> None:
        """
        Poprawia formatowanie list (punktowane i numerowane).

        Args:
            doc: Dokument python-docx
        """
        # Pattern dla list punktowanych: -, •, *, o
        bullet_pattern = re.compile(r'^[-•*o]\s+')

        # Pattern dla list numerowanych: a), (1), i., etc.
        numbered_pattern = re.compile(r'^(\([0-9a-z]+\)|[0-9a-z]+[\)\.:])\s+', re.IGNORECASE)

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Lista punktowana
            if bullet_pattern.match(text):
                paragraph.style = 'List Bullet'

            # Lista numerowana
            elif numbered_pattern.match(text):
                paragraph.style = 'List Number'
