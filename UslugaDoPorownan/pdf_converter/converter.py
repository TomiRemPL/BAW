"""
Główny moduł konwersji PDF→DOCX.

Dwupoziomowy system:
1. pdf2docx (podstawowy) - ~95% przypadków
2. pdfplumber (fallback) - ~5% skomplikowane tabele
"""

import logging
from pathlib import Path
from typing import Optional, Union, Tuple
from dataclasses import dataclass
import time

from pdf2docx import Converter as PDF2DOCXConverter
from docx import Document

from .config import PDFConverterConfig, DEFAULT_CONFIG
from .exceptions import ConversionError, TimeoutError, UnsupportedFormatError
from .validators import QualityValidator


logger = logging.getLogger(__name__)


@dataclass
class ConversionResult:
    """Wynik konwersji PDF→DOCX."""

    success: bool
    output_path: Optional[Path]
    quality_score: float  # 0.0-1.0
    conversion_time: float  # sekundy
    method: str  # "pdf2docx" lub "pdfplumber"
    error: Optional[str] = None
    fallback_used: bool = False

    def __str__(self) -> str:
        if self.success:
            return (
                f"Konwersja zakończona sukcesem (metoda: {self.method}, "
                f"jakość: {self.quality_score:.2f}, czas: {self.conversion_time:.2f}s)"
            )
        return f"Konwersja nie powiodła się: {self.error}"


class PDFConverter:
    """
    Konwerter PDF→DOCX z automatycznym fallback i walidacją jakości.

    Przykład użycia:
        >>> converter = PDFConverter()
        >>> result = converter.convert("input.pdf", "output.docx")
        >>> print(f"Jakość: {result.quality_score}")
    """

    def __init__(self, config: Optional[PDFConverterConfig] = None):
        """
        Inicjalizuje konwerter.

        Args:
            config: Opcjonalna konfiguracja. Jeśli None, używa DEFAULT_CONFIG.
        """
        self.config = config or DEFAULT_CONFIG
        self.validator = QualityValidator()

        logger.info(f"PDFConverter zainicjalizowany z konfiguracją: {self.config}")

    def convert(
        self,
        pdf_path: Union[str, Path],
        output_path: Union[str, Path],
        force_method: Optional[str] = None
    ) -> ConversionResult:
        """
        Konwertuje PDF na DOCX z automatycznym fallback.

        Args:
            pdf_path: Ścieżka do pliku PDF
            output_path: Ścieżka do wyjściowego pliku DOCX
            force_method: Wymuszenie metody ("pdf2docx" lub "pdfplumber"), domyślnie None (auto)

        Returns:
            ConversionResult z detalami konwersji

        Raises:
            ConversionError: Gdy konwersja się nie powiedzie
            TimeoutError: Gdy przekroczono limit czasu
            UnsupportedFormatError: Gdy plik nie jest PDF
        """
        pdf_path = Path(pdf_path)
        output_path = Path(output_path)

        # Walidacja wejścia
        self._validate_input(pdf_path)

        # Rozpocznij konwersję
        start_time = time.time()

        try:
            if force_method == "pdfplumber":
                return self._convert_with_pdfplumber(pdf_path, output_path, start_time)

            # Domyślnie próbuj pdf2docx
            result = self._convert_with_pdf2docx(pdf_path, output_path, start_time)

            # Sprawdź czy potrzebny fallback (niska jakość LUB błąd)
            if self.config.enable_fallback and force_method != "pdf2docx":
                should_fallback = False
                reason = ""

                if not result.success:
                    should_fallback = True
                    reason = "pdf2docx nie powiodło się"
                elif result.quality_score < self.config.min_quality_score:
                    should_fallback = True
                    reason = f"niska jakość ({result.quality_score:.2f})"

                if should_fallback:
                    logger.warning(
                        f"{reason}. Próba fallback do pdfplumber..."
                    )
                    return self._convert_with_pdfplumber(
                        pdf_path,
                        output_path,
                        start_time,
                        fallback=True
                    )

            return result

        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"Konwersja nie powiodła się: {e}")
            return ConversionResult(
                success=False,
                output_path=None,
                quality_score=0.0,
                conversion_time=elapsed,
                method="none",
                error=str(e)
            )

    def _validate_input(self, pdf_path: Path) -> None:
        """Waliduje plik wejściowy."""
        if not pdf_path.exists():
            raise ConversionError(f"Plik nie istnieje: {pdf_path}")

        if not pdf_path.is_file():
            raise ConversionError(f"Ścieżka nie jest plikiem: {pdf_path}")

        if pdf_path.suffix.lower() != ".pdf":
            raise UnsupportedFormatError(
                f"Nieobsługiwany format: {pdf_path.suffix}. Oczekiwano .pdf"
            )

    def _convert_with_pdf2docx(
        self,
        pdf_path: Path,
        output_path: Path,
        start_time: float
    ) -> ConversionResult:
        """
        Konwersja przy użyciu pdf2docx (podstawowa metoda).

        Args:
            pdf_path: Ścieżka do PDF
            output_path: Ścieżka wyjściowa DOCX
            start_time: Czas rozpoczęcia konwersji

        Returns:
            ConversionResult
        """
        logger.info(f"Konwersja pdf2docx: {pdf_path} -> {output_path}")

        try:
            # Konwersja
            cv = PDF2DOCXConverter(str(pdf_path))
            cv.convert(str(output_path))
            cv.close()

            elapsed = time.time() - start_time

            # Sprawdź timeout
            if elapsed > self.config.max_conversion_time:
                logger.warning(
                    f"Konwersja przekroczyła limit czasu: {elapsed:.2f}s > "
                    f"{self.config.max_conversion_time}s"
                )

            # Walidacja jakości
            quality_score = self.validator.validate(output_path, pdf_path)

            logger.info(
                f"Konwersja pdf2docx zakończona: {elapsed:.2f}s, "
                f"jakość: {quality_score:.2f}"
            )

            return ConversionResult(
                success=True,
                output_path=output_path,
                quality_score=quality_score,
                conversion_time=elapsed,
                method="pdf2docx",
                fallback_used=False
            )

        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"Błąd pdf2docx: {e}")

            # Zwróć wynik z błędem zamiast rzucać wyjątek
            # To pozwoli na automatyczny fallback do pdfplumber
            return ConversionResult(
                success=False,
                output_path=None,
                quality_score=0.0,
                conversion_time=elapsed,
                method="pdf2docx",
                error=f"pdf2docx nie powiodło się: {e}",
                fallback_used=False
            )

    def _convert_with_pdfplumber(
        self,
        pdf_path: Path,
        output_path: Path,
        start_time: float,
        fallback: bool = False
    ) -> ConversionResult:
        """
        Konwersja przy użyciu pdfplumber (fallback dla skomplikowanych tabel).

        Args:
            pdf_path: Ścieżka do PDF
            output_path: Ścieżka wyjściowa DOCX
            start_time: Czas rozpoczęcia konwersji
            fallback: Czy to fallback po pdf2docx

        Returns:
            ConversionResult
        """
        logger.info(f"Konwersja pdfplumber: {pdf_path} -> {output_path}")

        try:
            import pdfplumber
            from docx import Document
            from docx.shared import Pt, RGBColor

            # Stwórz nowy dokument DOCX
            doc = Document()

            # Otwórz PDF
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    if self.config.verbose:
                        logger.debug(f"Przetwarzanie strony {page_num}/{len(pdf.pages)}")

                    # Sprawdź timeout
                    elapsed = time.time() - start_time
                    if elapsed > self.config.max_conversion_time:
                        raise TimeoutError(
                            f"Przekroczono limit czasu: {elapsed:.2f}s"
                        )

                    # Wyodrębnij tabele (priorytet)
                    if self.config.preserve_tables:
                        tables = page.extract_tables()
                        for table in tables:
                            self._add_table_to_docx(doc, table)

                    # Wyodrębnij tekst
                    text = page.extract_text()
                    if text and text.strip():
                        doc.add_paragraph(text)

            # Zapisz dokument
            doc.save(str(output_path))

            elapsed = time.time() - start_time

            # Walidacja jakości
            quality_score = self.validator.validate(output_path, pdf_path)

            logger.info(
                f"Konwersja pdfplumber zakończona: {elapsed:.2f}s, "
                f"jakość: {quality_score:.2f}"
            )

            return ConversionResult(
                success=True,
                output_path=output_path,
                quality_score=quality_score,
                conversion_time=elapsed,
                method="pdfplumber",
                fallback_used=fallback
            )

        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"Błąd pdfplumber: {e}")
            raise ConversionError(f"pdfplumber nie powiodło się: {e}")

    def _add_table_to_docx(self, doc: Document, table_data: list) -> None:
        """
        Dodaje tabelę do dokumentu DOCX.

        Args:
            doc: Dokument python-docx
            table_data: Lista list z danymi tabeli
        """
        if not table_data:
            return

        # Znajdź wymiary tabeli
        rows = len(table_data)
        cols = max(len(row) for row in table_data) if table_data else 0

        if rows == 0 or cols == 0:
            return

        # Dodaj tabelę
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        # Wypełnij danymi
        for i, row_data in enumerate(table_data):
            for j, cell_value in enumerate(row_data):
                if j < cols:
                    cell = table.rows[i].cells[j]
                    cell.text = str(cell_value) if cell_value else ""

    def convert_batch(
        self,
        pdf_paths: list[Union[str, Path]],
        output_dir: Union[str, Path]
    ) -> list[ConversionResult]:
        """
        Konwertuje wiele plików PDF.

        Args:
            pdf_paths: Lista ścieżek do plików PDF
            output_dir: Katalog wyjściowy dla plików DOCX

        Returns:
            Lista ConversionResult dla każdego pliku
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        results = []
        for pdf_path in pdf_paths:
            pdf_path = Path(pdf_path)
            output_path = output_dir / f"{pdf_path.stem}.docx"

            try:
                result = self.convert(pdf_path, output_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Błąd konwersji {pdf_path}: {e}")
                results.append(ConversionResult(
                    success=False,
                    output_path=None,
                    quality_score=0.0,
                    conversion_time=0.0,
                    method="none",
                    error=str(e)
                ))

        return results
