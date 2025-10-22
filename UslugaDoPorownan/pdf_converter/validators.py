"""
Walidatory jakości konwersji PDF→DOCX.

System walidacji zwraca wynik 0.0-1.0 bazujący na:
- Liczbie wyodrębnionych paragrafów
- Liczbie wyodrębnionych tabel
- Strukturze dokumentu
- Formatowaniu
"""

import logging
from pathlib import Path
from typing import Optional
from docx import Document
import pdfplumber


logger = logging.getLogger(__name__)


class QualityValidator:
    """
    Walidator jakości konwersji PDF→DOCX.

    Zwraca wynik 0.0-1.0 gdzie:
    - 0.9-1.0: Doskonała konwersja
    - 0.7-0.9: Dobra konwersja
    - 0.5-0.7: Średnia konwersja (rozważ fallback)
    - 0.0-0.5: Słaba konwersja (użyj fallback)
    """

    def validate(
        self,
        docx_path: Path,
        pdf_path: Optional[Path] = None
    ) -> float:
        """
        Waliduje jakość konwersji.

        Args:
            docx_path: Ścieżka do wygenerowanego DOCX
            pdf_path: Opcjonalna ścieżka do oryginalnego PDF (dla porównania)

        Returns:
            Wynik jakości 0.0-1.0
        """
        try:
            # Otwórz DOCX
            doc = Document(docx_path)

            scores = []

            # 1. Sprawdź paragrafy (waga 30%)
            paragraph_score = self._validate_paragraphs(doc)
            scores.append(("paragraphs", paragraph_score, 0.30))

            # 2. Sprawdź tabele (waga 40% - priorytet!)
            table_score = self._validate_tables(doc)
            scores.append(("tables", table_score, 0.40))

            # 3. Sprawdź strukturę (waga 20%)
            structure_score = self._validate_structure(doc)
            scores.append(("structure", structure_score, 0.20))

            # 4. Sprawdź formatowanie (waga 10%)
            formatting_score = self._validate_formatting(doc)
            scores.append(("formatting", formatting_score, 0.10))

            # Jeśli mamy PDF, porównaj z oryginałem
            if pdf_path and pdf_path.exists():
                comparison_score = self._compare_with_pdf(doc, pdf_path)
                scores.append(("comparison", comparison_score, 0.20))
                # Normalizuj wagi
                total_weight = sum(weight for _, _, weight in scores)
                scores = [(name, score, weight / total_weight) for name, score, weight in scores]

            # Oblicz ważoną średnią
            final_score = sum(score * weight for _, score, weight in scores)

            # Log szczegółów
            logger.info(
                f"Wynik walidacji dla {docx_path.name}: {final_score:.2f} "
                f"(paragrafy: {paragraph_score:.2f}, tabele: {table_score:.2f}, "
                f"struktura: {structure_score:.2f}, formatowanie: {formatting_score:.2f})"
            )

            return round(final_score, 2)

        except Exception as e:
            logger.error(f"Błąd walidacji: {e}")
            return 0.0

    def _validate_paragraphs(self, doc: Document) -> float:
        """
        Waliduje paragrafy dokumentu.

        Args:
            doc: Dokument python-docx

        Returns:
            Wynik 0.0-1.0
        """
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]

        if not paragraphs:
            return 0.0

        # Sprawdź czy są niepuste paragrafy
        non_empty_count = len(paragraphs)

        # Heurystyka: dokument powinien mieć przynajmniej kilka paragrafów
        if non_empty_count == 0:
            return 0.0
        elif non_empty_count < 5:
            return 0.5
        elif non_empty_count < 10:
            return 0.7
        else:
            return 1.0

    def _validate_tables(self, doc: Document) -> float:
        """
        Waliduje tabele dokumentu (priorytet!).

        Args:
            doc: Dokument python-docx

        Returns:
            Wynik 0.0-1.0
        """
        tables = doc.tables

        if not tables:
            # Brak tabel - może to być OK jeśli dokument nie ma tabel
            # Przyjmujemy neutralny wynik
            return 0.8

        # Sprawdź jakość tabel
        valid_tables = 0
        for table in tables:
            if self._is_valid_table(table):
                valid_tables += 1

        if len(tables) == 0:
            return 0.8  # Brak tabel to nie błąd

        # Procent prawidłowych tabel
        ratio = valid_tables / len(tables)
        return ratio

    def _is_valid_table(self, table) -> bool:
        """
        Sprawdza czy tabela jest prawidłowo wyodrębniona.

        Args:
            table: Tabela python-docx

        Returns:
            True jeśli tabela jest OK
        """
        # Tabela musi mieć przynajmniej 1 wiersz i 1 kolumnę
        if len(table.rows) < 1 or len(table.columns) < 1:
            return False

        # Sprawdź czy przynajmniej połowa komórek ma treść
        total_cells = len(table.rows) * len(table.columns)
        non_empty_cells = 0

        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    non_empty_cells += 1

        return non_empty_cells / total_cells >= 0.3  # Przynajmniej 30% komórek ma treść

    def _validate_structure(self, doc: Document) -> float:
        """
        Waliduje strukturę dokumentu (nagłówki, numeracja).

        Args:
            doc: Dokument python-docx

        Returns:
            Wynik 0.0-1.0
        """
        score = 0.0

        # Sprawdź nagłówki
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        if headings:
            score += 0.5

        # Sprawdź listy numerowane/punktowane
        lists = [p for p in doc.paragraphs if p.style.name.startswith('List')]
        if lists:
            score += 0.5

        return min(score, 1.0)

    def _validate_formatting(self, doc: Document) -> float:
        """
        Waliduje formatowanie dokumentu (czcionki, pogrubienia).

        Args:
            doc: Dokument python-docx

        Returns:
            Wynik 0.0-1.0
        """
        score = 0.0

        # Sprawdź czy zachowano formatowanie (bold, italic)
        formatted_runs = 0
        total_runs = 0

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                total_runs += 1
                if run.bold or run.italic or run.underline:
                    formatted_runs += 1

        if total_runs > 0 and formatted_runs > 0:
            score = 1.0
        elif total_runs > 0:
            score = 0.5
        else:
            score = 0.8  # Brak runs nie oznacza błędu

        return score

    def _compare_with_pdf(self, doc: Document, pdf_path: Path) -> float:
        """
        Porównuje DOCX z oryginalnym PDF.

        Args:
            doc: Dokument python-docx
            pdf_path: Ścieżka do PDF

        Returns:
            Wynik 0.0-1.0
        """
        try:
            with pdfplumber.open(pdf_path) as pdf:
                # Porównaj liczbę stron
                pdf_pages = len(pdf.pages)

                # W DOCX nie ma bezpośredniego API do stron, więc używamy heurystyki
                # Zakładamy ~3 paragrafy na stronę
                estimated_docx_pages = len([p for p in doc.paragraphs if p.text.strip()]) / 3

                if pdf_pages == 0:
                    return 0.8

                # Porównaj liczbę tabel
                pdf_tables_count = sum(len(page.extract_tables()) for page in pdf.pages)
                docx_tables_count = len(doc.tables)

                tables_ratio = 1.0
                if pdf_tables_count > 0:
                    tables_ratio = min(docx_tables_count / pdf_tables_count, 1.0)

                # Średnia z porównań
                return tables_ratio

        except Exception as e:
            logger.warning(f"Nie można porównać z PDF: {e}")
            return 0.8  # Neutralny wynik przy błędzie
